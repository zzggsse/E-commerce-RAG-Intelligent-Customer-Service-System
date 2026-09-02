"""文档解析、清洗、切片。

两类文档差异化处理：
  goods    —— 商品资料，参数密集，小切片 + 按行边界切，避免一条参数被切散
  aftersale—— 售后 FAQ，先按"问答对边界"切成完整 QA，再对超长 QA 二次切分
"""
import logging
import re
from pathlib import Path
from typing import Dict, List

import settings

logger = logging.getLogger(__name__)

SUPPORTED_SUFFIX = {".md", ".markdown", ".txt", ".pdf", ".docx", ".csv", ".tsv"}

# 需要清洗掉的广告 / 无效内容
_AD_PATTERNS = [
    r"^\s*[-*=~_]{3,}\s*$",
    r"关注(店铺|本店).{0,20}(优惠|领券|福利)",
    r"(点击|长按).{0,10}(领取|购买|下单)",
    r"扫码加(微信|好友)",
    r"^\s*第\s*\d+\s*页\s*$",
    r"^\s*Page\s*\d+\s*$",
]
_AD_RE = [re.compile(p) for p in _AD_PATTERNS]
_FAQ_RE = [re.compile(p) for p in settings.FAQ_QUESTION_PATTERNS]


# ---------------------------------------------------------------- 读取
def read_file(path: str) -> str:
    suffix = Path(path).suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix == ".csv":
        return _read_csv(path)
    if suffix == ".tsv":
        return _read_tsv(path)
    return Path(path).read_text(encoding="utf-8", errors="ignore")


def _read_pdf(path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _read_docx(path: str) -> str:
    from docx import Document

    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts).strip()


def _read_csv(path: str) -> str:
    return _read_delimited(path, ",")


def _read_tsv(path: str) -> str:
    return _read_delimited(path, "\t")


def _read_delimited(path: str, delimiter: str) -> str:
    """CSV/TSV: 每行转成 列名：值 的结构化文本，便于按行切块检索。"""
    import csv

    with open(path, encoding="utf-8-sig", errors="ignore", newline="") as fp:
        rows = list(csv.reader(fp, delimiter=delimiter))
    if not rows:
        return ""

    def is_header(row):
        return bool(row) and not any(_is_numeric(c) for c in row[: min(3, len(row))])

    header = rows[0] if is_header(rows[0]) else None
    body = rows[1:] if header else rows
    out = []
    for row in body:
        cells = [c.strip() for c in row]
        if not any(cells):
            continue
        if header:
            pairs = []
            for i, cell in enumerate(cells):
                name = (header[i].strip() if i < len(header) and header[i].strip() else "字段%d" % (i + 1))
                pairs.append("%s：%s" % (name, cell))
            out.append("，".join(pairs))
        else:
            out.append(" | ".join(cells))
    return "\n".join(out).strip()


def _is_numeric(s: str) -> bool:
    t = s.strip().replace(",", "").replace("%", "").replace("¥", "").replace("元", "")
    try:
        float(t)
        return True
    except ValueError:
        return False


# ---------------------------------------------------------------- 清洗
def clean_text(text: str) -> str:
    """去广告、去乱码、压缩空行，保留 Markdown 结构。"""
    text = text.lstrip("\ufeff")
    lines = []
    for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = line.replace("\u3000", " ").rstrip()
        if any(r.search(line) for r in _AD_RE):
            continue
        # 去掉不可见控制字符
        line = re.sub(r"[\x00-\x08\x0b-\x1f\x7f]", "", line)
        # 乱码行过滤：非中英数标点占比过高
        stripped = line.strip()
        if len(stripped) >= 8:
            valid = len(re.findall(r"[\u4e00-\u9fa5A-Za-z0-9\s\W]", stripped))
            if valid / len(stripped) < 0.6:
                continue
        lines.append(line)
    text = "\n".join(lines)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ---------------------------------------------------------------- 切片
def split_document(text: str, doc_type: str) -> List[str]:
    if doc_type == "aftersale":
        return _split_faq(text)
    return _split_goods(text)


def _split_goods(text: str) -> List[str]:
    """商品资料切片：以行为最小单位，绝不在一行（一条参数）中间断开。"""
    cfg = settings.CHUNK_CONFIG["goods"]
    return _pack_by_lines(text, cfg["chunk_size"], cfg["chunk_overlap"])


def _split_faq(text: str) -> List[str]:
    """售后 FAQ 切片：先按问答对边界切，保证"问题+答案"永远在同一片段。"""
    cfg = settings.CHUNK_CONFIG["aftersale"]
    size, overlap = cfg["chunk_size"], cfg["chunk_overlap"]

    blocks: List[str] = []
    current: List[str] = []
    for line in text.split("\n"):
        if _is_question_line(line) and current:
            blocks.append("\n".join(current).strip())
            current = [line]
        else:
            current.append(line)
    if current:
        blocks.append("\n".join(current).strip())
    blocks = [b for b in blocks if b]

    if len(blocks) <= 1:  # 没识别出问答结构，退化为按行打包
        return _pack_by_lines(text, size, overlap)

    chunks: List[str] = []
    buffer = ""
    for block in blocks:
        if len(block) > size:  # 单个 QA 超长，单独二次切分
            if buffer:
                chunks.append(buffer)
                buffer = ""
            chunks.extend(_pack_by_lines(block, size, overlap))
            continue
        if not buffer:
            buffer = block
        elif len(buffer) + len(block) + 2 <= size:
            buffer = buffer + "\n\n" + block  # 小 QA 合并，避免碎片
        else:
            chunks.append(buffer)
            buffer = block
    if buffer:
        chunks.append(buffer)
    return chunks


def _is_question_line(line: str) -> bool:
    return any(r.search(line) for r in _FAQ_RE)


def _pack_by_lines(text: str, size: int, overlap: int) -> List[str]:
    """按行贪心打包成不超过 size 的片段，相邻片段回溯 overlap 字符做重叠。"""
    units: List[str] = []
    for line in text.split("\n"):
        if len(line) <= size:
            units.append(line)
        else:  # 超长行按句子再切
            units.extend(_split_long_line(line, size))

    chunks: List[str] = []
    buffer = ""
    for unit in units:
        candidate = unit if not buffer else buffer + "\n" + unit
        if len(candidate) <= size:
            buffer = candidate
            continue
        if buffer:
            chunks.append(buffer.strip())
            tail = buffer[-overlap:] if overlap else ""
            buffer = (tail + "\n" + unit).strip() if tail else unit
        else:
            buffer = unit
    if buffer.strip():
        chunks.append(buffer.strip())
    return [c for c in chunks if c]


def _split_long_line(line: str, size: int) -> List[str]:
    parts = re.split(r"(?<=[。！？；;!?])", line)
    out, buf = [], ""
    for part in parts:
        if len(buf) + len(part) <= size:
            buf += part
        else:
            if buf:
                out.append(buf)
            while len(part) > size:
                out.append(part[:size])
                part = part[size:]
            buf = part
    if buf:
        out.append(buf)
    return out


# ---------------------------------------------------------------- 主流程
def parse_document(
    path: str,
    doc_type: str,
    goods_id: str = "",
    category: str = "",
    source: str = "",
) -> List[Dict]:
    """解析单个文件，返回带元数据的切片列表。"""
    suffix = Path(path).suffix.lower()
    if suffix not in SUPPORTED_SUFFIX:
        raise ValueError(f"不支持的文件类型: {suffix}")

    raw = read_file(path)
    text = clean_text(raw)
    if not text:
        return []

    chunks = split_document(text, doc_type)
    source = source or Path(path).name
    result = []
    for idx, chunk in enumerate(chunks):
        result.append(
            {
                "content": chunk,
                "doc_type": doc_type,
                # 通用售后文档 goods_id 留空字符串，表示对所有商品生效
                "goods_id": goods_id or "",
                "category": category or "",
                "source": source,
                "chunk_index": idx,
            }
        )
    logger.info("解析 %s -> %d 个切片 (doc_type=%s)", source, len(result), doc_type)
    return result
